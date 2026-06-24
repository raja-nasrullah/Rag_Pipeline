# RAG_Pipeline.py
import os
import chromadb
from openai import OpenAI
from PyPDF2 import PdfReader
from docx import Document

client = OpenAI(api_key="Key")  # Replace with your OpenAI API key
chroma_client = chromadb.PersistentClient(path="chroma_db")
collection = chroma_client.get_or_create_collection(name="file_embeddings")

# ----------- TEXT EXTRACTION -----------
def extract_text_from_file(file_path):
    text = ""
    if file_path.endswith(".txt"):
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
    elif file_path.endswith(".pdf"):
        reader = PdfReader(file_path)
        for page in reader.pages:
            text += page.extract_text() or ""
    elif file_path.endswith(".docx"):
        doc = Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
    return text.strip()

def chunk_text(text, max_chars=2000):
    for i in range(0, len(text), max_chars):
        yield text[i:i + max_chars]

# ----------- EMBEDDINGS -----------
# ------------------- EMBEDDING PIPELINE -------------------

def extract_text_from_file(file_path: str) -> str:
    """
    Extract readable text from TXT, PDF, DOCX files.
    """
    text = ""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".txt":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()

    elif ext == ".pdf":
        try:
            reader = PdfReader(file_path)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except Exception as e:
            print(f"⚠️ PDF extraction failed: {e}")

    elif ext == ".docx":
        try:
            doc = Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        except Exception as e:
            print(f"⚠️ DOCX extraction failed: {e}")

    else:
        print(f"⚠️ Unsupported file type: {ext}")

    return text.strip()

def split_text_into_chunks(text, chunk_size=500, overlap=50):
    """
    Split text into smaller chunks for embeddings.
    """
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start = end - overlap  # allow overlap
    return chunks


def generate_embeddings(chunks, model="text-embedding-3-small", batch_size=100):
    """
    Generate embeddings for a list of text chunks in smaller batches.
    Avoids hitting the max token/request limit.
    """
    all_embeddings = []
    for i in range(0, len(chunks), batch_size):
        batch = chunks[i:i+batch_size]
        response = client.embeddings.create(
            model=model,
            input=batch
        )
        embeddings = [item.embedding for item in response.data]
        all_embeddings.extend(embeddings)
        print(f"⚡ Embedded {i+len(batch)}/{len(chunks)} chunks")
    return all_embeddings


# ----------- STORAGE (new) -----------
def store_embeddings(file_name: str, chunks: list[str], embeddings: list[list[float]], batch_size: int = 200):
    """Store chunks + embeddings in ChromaDB in smaller batches."""
    total = len(chunks)
    for i in range(0, total, batch_size):
        batch_chunks = chunks[i:i+batch_size]
        batch_embeddings = embeddings[i:i+batch_size]
        batch_ids = [f"{file_name}_part{j}" for j in range(i, i+len(batch_chunks))]

        collection.add(
            ids=batch_ids,
            documents=batch_chunks,
            metadatas=[{"source": file_name}] * len(batch_chunks),
            embeddings=batch_embeddings
        )
        print(f"💾 Stored {i+len(batch_chunks)}/{total} chunks for {file_name}")

    print(f"✅ Finished storing {total} chunks for {file_name}")

# ----------- DELETE EMBEDDINGS -----------


# def delete_local_file(file_id, file_map):
#     filename = file_map.get(file_id)
#     if filename:
#         local_path = os.path.join("data", filename)

#         if os.path.exists(local_path):
#             os.remove(local_path)
#             print(f"🗑️ Deleted local file: {local_path}")
#         else:
#             print(f"⚠️ Local file not found: {local_path}")


#         # Update file map
#         file_map.pop(file_id, None)
#         save_file_map(file_map)
#     else:
#         print(f"⚠️ No mapping found for deleted fileId: {file_id}")


def embed_and_store(file_path: str):
    file_name = os.path.basename(file_path)
    print(f"📂 Processing file for embeddings: {file_name}")

    text = extract_text_from_file(file_path)
    print(f"📝 Extracted {len(text)} characters from {file_name}")

    chunks = list(chunk_text(text))
    print(f"✂️ Split into {len(chunks)} chunks")

    embeddings = generate_embeddings(chunks)   # ✅ real embeddings now
    print(f"🧠 Generated {len(embeddings)} embeddings")

    store_embeddings(file_name, chunks, embeddings)
    print(f"🎉 Embedding pipeline complete for {file_name}")

# ----------- QUERY / CHATBOT -----------
def chatbot():
    print("🤖 Chatbot started (type 'exit' to quit')")
    while True:
        q = input("You: ")
        if q.lower() in ["exit", "quit"]:
            break

        # Step 1: embed query
        query_embedding = client.embeddings.create(
            model="text-embedding-3-small",
            input=q
        ).data[0].embedding

        # Step 2: query Chroma
        results = collection.query(
            query_embeddings=[query_embedding],
            n_results=5,
            include=["documents", "distances"]
        )

        if not results["documents"] or not results["documents"][0]:
            print("Bot: I couldn’t find anything in your documents.")
            continue

        # Step 3: clean up docs
        docs = results["documents"][0][:9]
        valid_docs = [d for d in docs if d]  # remove None

        if not valid_docs:
            print("Bot: I don’t know (no valid documents found).")
            continue

        context_text = "\n".join(valid_docs)

        # Step 4: Ask GPT
        try:
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": (
                        "You are a strict assistant. "
                        "Answer ONLY based on the provided document context. "
                        "If the answer is not in the context, reply with: "
                        "'I don’t know based on the documents.'"
                    )},
                    {"role": "user", "content": f"Context:\n{context_text}\n\nQuestion: {q}"}
                ]
            )
            print("Bot:", response.choices[0].message.content.strip())
        except Exception as e:
            print(f"⚠️ Error answering query: {e}")


# ------------------- MAIN -------------------
if __name__ == "__main__":
    chatbot()
